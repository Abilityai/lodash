import docx
import magic

from io import BytesIO
from PyPDF2 import PdfReader
import pymupdf
import aiohttp
from datetime import datetime
import os
from io import BytesIO
import fitz
from typing import Tuple, List
from urllib.parse import urlparse
from ftfy.fixes import remove_control_chars
from docx import Document
import uuid
import io

def read_file_content(file: BytesIO):
    mime = magic.Magic(mime=True)
    file_type = mime.from_buffer(file.getvalue())

    extracted_text = None
    if 'pdf' in file_type.lower():
        reader = PdfReader(file)
        text = ''
        for page in reader.pages:
            text += page.extract_text() or ''
        extracted_text = text.strip() or "<Could not extract text from the PDF>"
    elif 'officedocument' in file_type.lower() or 'msword' in file_type.lower():
        doc = docx.Document(file)
        text = '\n'.join([paragraph.text for paragraph in doc.paragraphs])
        extracted_text = text.strip() or "<Could not extract text from the DOCX file>"

    return extracted_text, file_type

def pdf_to_png(file: BytesIO, max_height=2048, max_width=2048):
    mime = magic.Magic(mime=True)
    file_type = mime.from_buffer(file.getvalue())
    file_bytes = file.getvalue()

    result = []

    if 'pdf' in file_type.lower():
        pdf = pymupdf.open(stream=file_bytes, filetype="pdf")
        for page_num in range(pdf.page_count):
            page = pdf.load_page(page_num)
            rect = page.rect
            scale = min(max_width/rect.width, max_height/rect.height)
            matrix = pymupdf.Matrix(scale, scale)
            pixmap = page.get_pixmap(matrix=matrix)
            img = pixmap.tobytes("png")
            result.append(BytesIO(img))
        pdf.close()
    return result


async def process_document(url: str, upload_callback) -> dict:
    """Process a document from a URL and return a structured document. Upload_callback would be used to upload images to the server."""
    # Step 1: Download the file
    async with aiohttp.ClientSession() as session:
        async with session.get(url) as response:
            if response.status != 200:
                raise ValueError(f"Failed to download document from URL: {url}")
            file_bytes = await response.read()

    # Step 2: Determine the file type using python-magic
    mime = magic.Magic(mime=True)
    file_type = mime.from_buffer(file_bytes)

    # Check the file type
    if file_type not in ['application/pdf', 'application/vnd.openxmlformats-officedocument.wordprocessingml.document']:
        raise ValueError(f"Unsupported file type: {file_type}")

    # Initialize structures
    contents = []
    images = []
    metadata = {
        "filename": os.path.basename(urlparse(url).path),
        "created_at": datetime.utcnow().isoformat() + "Z"
    }

    if file_type == 'application/pdf':
        # Process PDF
        doc = fitz.open(stream=file_bytes, filetype="pdf")
        image_count = 0

        for page_num in range(len(doc)):
            page = doc.load_page(page_num)
            text = page.get_text("markdown")
            cleaned_text = remove_control_chars(text)
            contents.append(cleaned_text)

            # Extract images
            image_list = page.get_images(full=True)
            for img_index, img in enumerate(image_list):
                xref = img[0]
                base_image = doc.extract_image(xref)
                image_bytes = base_image["image"]
                image_ext = base_image["ext"]
                image_filename = f"{uuid.uuid4()}.{image_ext}"

                # Upload image using callback
                image_info = await upload_callback(filename=image_filename, file=io.BytesIO(image_bytes))
                images.append({
                    "type": "image",
                    "url": image_info["url"]
                })
                image_count += 1

    elif file_type == 'application/vnd.openxmlformats-officedocument.wordprocessingml.document':
        # Process DOCX

        doc = Document(io.BytesIO(file_bytes))
        image_count = 0

        for para in doc.paragraphs:
            cleaned_text = remove_control_chars(para.text)
            contents.append(cleaned_text)

        for rel in doc.part.rels.values():
            if "image" in rel.reltype:
                image = rel.target
                image_bytes = image.blob
                image_ext = image.content_type.split('/')[-1]
                image_filename = f"{uuid.uuid4()}.{image_ext}"

                # Upload image using callback
                image_info = await upload_callback(filename=image_filename, file=io.BytesIO(image_bytes))
                images.append({
                    "type": "image",
                    "url": image_info["url"]
                })
                image_count += 1

    # Convert contents to Markdown sections (if not already in Markdown)
    # Here, assuming contents are already in suitable Markdown format
    structured_document = {
        "type": "document",
        "contents": contents,
        "images": images,
        "metadata": metadata
    }

    return structured_document
